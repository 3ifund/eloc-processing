"""
Document text extraction classes for PDF and Word documents
"""
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Optional
import logging

logger = logging.getLogger(__name__)


class DocumentExtractor(ABC):
    
    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        pass
    
    @abstractmethod
    def supports(self, content_type: str) -> bool:
        """Check if this extractor supports the content type"""
        pass


class PDFExtractor(DocumentExtractor):
    """Extract text from PDF documents using pdfplumber"""
    
    def __init__(self):
        try:
            import pdfplumber
            self.pdfplumber = pdfplumber
            self.available = True
        except ImportError:
            logger.error("pdfplumber not installed - pip install pdfplumber")
            self.available = False
    
    def supports(self, content_type: str) -> bool:
        """Check if this is a PDF"""
        return content_type == "application/pdf"
    
    def extract_text(self, file_path: str, max_pages: int = 10) -> str:
        """Extract text from PDF file"""
        if not self.available:
            raise RuntimeError("pdfplumber is not installed")
        
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        logger.info(f"Extracting text from PDF: {file_path.name}")
        
        text_parts = []
        
        try:
            with self.pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                pages_to_process = min(total_pages, max_pages)
                
                logger.info(f"PDF has {total_pages} pages, processing first {pages_to_process}")
                
                for i, page in enumerate(pdf.pages[:pages_to_process]):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                extracted_text = "\n".join(text_parts)
                char_count = len(extracted_text)
                
                logger.info(f"Extracted {char_count} characters from {pages_to_process} pages")
                
                return extracted_text
                
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise


class WordExtractor(DocumentExtractor):
    """Extract text from Word documents (.docx)"""
    
    def __init__(self):
        try:
            from docx import Document
            self.Document = Document
            self.available = True
        except ImportError:
            logger.error("python-docx not installed - pip install python-docx")
            self.available = False
    
    def supports(self, content_type: str) -> bool:
        """Check if this is a Word document"""
        return content_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword"
        ]
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from Word document"""
        if not self.available:
            raise RuntimeError("python-docx is not installed")
        
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        logger.info(f"Extracting text from Word doc: {file_path.name}")
        
        try:
            doc = self.Document(file_path)
            
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            extracted_text = "\n".join(text_parts)
            char_count = len(extracted_text)
            
            logger.info(f"Extracted {char_count} characters from Word document")
            
            return extracted_text
            
        except Exception as e:
            logger.error(f"Error extracting Word text: {e}")
            raise


class ExtractorFactory:
    """Factory to get the right extractor for a document type"""
    
    def __init__(self):
        self.extractors = [
            PDFExtractor(),
            WordExtractor()
        ]
    
    def get_extractor(self, content_type: str) -> Optional[DocumentExtractor]:
        """Get the appropriate extractor for a content type"""
        for extractor in self.extractors:
            if extractor.supports(content_type):
                return extractor
        return None
    
    def extract_text(self, file_path: str, content_type: str) -> Optional[str]:
        """Extract text from a document"""
        extractor = self.get_extractor(content_type)
        
        if not extractor:
            logger.warning(f"No extractor available for content type: {content_type}")
            return None
        
        try:
            return extractor.extract_text(file_path)
        except Exception as e:
            logger.error(f"Extraction failed for {file_path}: {e}")
            return None